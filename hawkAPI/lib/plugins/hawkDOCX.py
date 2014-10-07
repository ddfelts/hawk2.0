from hawkAPI.lib.core.hawklib import hawklib
from docx import Document
from docx.shared import Inches

class hawkDOCX():

    def __init__(self,name,template,hawk):
        self.doc = Document(template)
        self.secs = self.doc.sections
        self.name = name
        self.hlib = hawklib(hawk)
        self.cur = None

    def addPara(self,data):
        self.cur = self.doc.add_paragraph(data)

    def addHead(self,data):
        self.doc.add_heading(data,level=1)

    def addPageBreak(self):
        self.doc.add_page_break()

    def addImage(self,image,w=None):
        if w == None:
            self.doc.add_picture(image)
        else:
            self.doc.add_picture(image,width=Inches(w))

    def addTable(self,data,nkeys=None):
        if not nkeys:
            keys = self.hlib.getKeys(data)
        else:
            keys = nkeys
        table = self.doc.add_table(rows=1,cols=len(keys))
        row = table.rows[0]
        for i in range(len(keys)):
            row.cells[i].text = str(keys[i])
        for x in data:
            nrow = table.add_row().cells
            for b in range(len(keys)):
                if keys[b] not in x:
                    outb = None
                else:
                    outb = x[keys[b]]
                nrow[b].text = str(outb)

    def saveDoc(self):
        self.doc.save(self.name)
